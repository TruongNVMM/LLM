#!/usr/bin/env python3
"""
pipeline.py – End-to-end HUST RAG Pipeline: Crawl → Clean → Chunk → PostgreSQL

Luồng xử lý:
    1. Gọi API HUST (GetWebTitleLst + GetWebTitleInfo)
    2. Với mỗi bài viết thô:
       a. normalize_article()   → làm sạch HTML, extract text, links
       b. enrich_local_files()  → đọc file .doc/.docx/.pdf từ data_fetch/ (tuỳ chọn)
       c. to_rag_document()     → build RAG document format
       d. build_chunks()        → cắt đoạn thành chunks
       e. upsert_article()      → lưu bài đã làm sạch vào PostgreSQL
       f. upsert_chunks()       → lưu chunks vào PostgreSQL
    3. Ghi log pipeline_run vào PostgreSQL

Dữ liệu được lưu trực tiếp vào PostgreSQL sau khi xử lý.
Không còn file JSONL/JSON trung gian.

Sử dụng:
    python src/data_processing/pipeline.py
    python src/data_processing/pipeline.py --no-detail
    python src/data_processing/pipeline.py --fetch-attachments
    python src/data_processing/pipeline.py --enrich-local
    python src/data_processing/pipeline.py --enrich-local --data-fetch-dir path/to/data_fetch
    python src/data_processing/pipeline.py --verbose
    python src/data_processing/pipeline.py --init-schema-only

Yêu cầu:
    pip install psycopg2-binary pypdf python-docx markdownify

Cấu hình (file .env hoặc biến môi trường):
    POSTGRES_HOST, POSTGRES_PORT, POSTGRES_DB, POSTGRES_USER, POSTGRES_PASSWORD
"""

from __future__ import annotations

import argparse
import hashlib
import io
import json
import logging
import os
import re
import sys
import time
import unicodedata
import warnings
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

# ---------------------------------------------------------------------------
# Encoding fix cho Windows
# ---------------------------------------------------------------------------
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8")

# ---------------------------------------------------------------------------
# Path setup: cho phép import từ thư mục cha (data_processing/)
# ---------------------------------------------------------------------------
_THIS_DIR = Path(__file__).parent.resolve()
if str(_THIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THIS_DIR))

# ---------------------------------------------------------------------------
# Import functions từ crawler hiện có (tái sử dụng toàn bộ logic)
# ---------------------------------------------------------------------------
from crawl_hust_sotay import (  # noqa: E402
    API_BASE,
    PAGE_BASE,
    CATEGORIES,
    TYPE_DOC_TO_CATEGORY,
    ATTACHMENT_FETCH_DELAY,
    normalize_article,
    to_rag_document,
    post_json,
    enrich_links,
    clean_text,
)

# ---------------------------------------------------------------------------
# Import chunking functions (tái sử dụng logic chunking)
# ---------------------------------------------------------------------------
from chunking_pipeline import (  # noqa: E402
    build_chunks,
    CHUNK_SIZE,
    CHUNK_OVERLAP,
)

# ---------------------------------------------------------------------------
# Import convert script runner & DB layer
# ---------------------------------------------------------------------------
from integrate_data_fetch import run_convert_script  # noqa: E402

from db.connection import (  # noqa: E402
    get_managed_connection,
    init_schema,
    upsert_article,
    upsert_chunks,
    create_pipeline_run,
    finish_pipeline_run,
)

# ---------------------------------------------------------------------------
# Optional imports cho local file extraction
# ---------------------------------------------------------------------------
try:
    from pypdf import PdfReader as _PdfReader
    logging.getLogger("pypdf").setLevel(logging.ERROR)
except ImportError:
    try:
        from PyPDF2 import PdfReader as _PdfReader  # type: ignore[no-redef]
        logging.getLogger("PyPDF2").setLevel(logging.ERROR)
    except ImportError:
        _PdfReader = None  # type: ignore[assignment]

try:
    import docx as _docx
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

try:
    import win32com.client as _win32com
    import pythoncom as _pythoncom
    _HAS_WIN32 = True
except ImportError:
    _HAS_WIN32 = False

# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%H:%M:%S",
)
logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Cấu hình
# ---------------------------------------------------------------------------
_THIS_DIR = Path(__file__).resolve().parent
DEFAULT_DATA_FETCH_DIR = _THIS_DIR / "data_fetch"

# Mapping: tên file → (doc_id, chữ ký URL để match với link trong bài)
# Lấy từ integrate_data_fetch.py, cập nhật tại đây khi có file mới
FILE_MAPPING: list[tuple[str, int, str]] = [
    ("16. Đơn công nhận HP tương đương.doc",                  69,  "EUZMK0oVAZpB"),
    ("20. Miễn ngoại ngữ.docx",                               69,  "EVM4qdnP_45J"),
    ("HƯỚNG DẪN VỀ VIỆC YÊU CẦU NHẬP ĐIỂM ĐÚNG THỜI HẠN.docx", 69, "EUDPbbk1m01J"),
    ("08.2 Đơn thắc mắc ĐATN.doc",                            69,  "EdDCJnCWLlRH"),
    ("22. Mẫu đơn rút học phần.doc",                          69,  "Eap5qh9aMf5E"),
    ("03.4 Đơn xin thôi học_ĐT.QT06.BM.01_Rev 16.4.2023.doc", 69, "EdHS7TFGfW1G"),
    ("04.3 Đơn xin nghỉ dài hạn_06.02 ĐT.QT06.BM.02.doc",    69,  "EVDUxHPtC3NP"),
    ("17.1 Đơn xin chuyển trường ĐH-Rev 01_Updated 13.4.2023.doc", 69, "ETUYBfTNy3pA"),
    ("13. Mẫu đơn đăng ký sang hệ VHVL.docx",                69,  "EUyvO31qrjRD"),
    ("19. Trích sao bảng điểm.docx",                          69,  "EeCW4MRDLJZN"),
    ("Hướng dẫn hoãn thi-thi bù.docx",                       69,  "EVmtZeRjgWFH"),
    ("15.1 Đơn hoãn thi_07.01 DT.QT.07.BM.01_v2.doc",        69,  "ESj7AYgEYQZM"),
    ("Đơn xin chuyển ngành học KSCS.doc",                     69,  "ERR1zxfXQIFO"),
    ("20230710 1. QĐ Học bổng KKHT 2023.pdf",                 61,  "ESCBPVQlzNFOlglsvRwxAZYBT5EgEPuy"),
    ("Quy định Học bổng Trần Đại Nghĩa 2025.docx",            61,  "ERczoBK8xr9P"),
    ("Quy định xét cấp HB tài trợ 2024 LasVer.pdf",           61,  "EbEKhfFyCe9C"),
    ("QĐ HB gắn kết quê hương 02.4.2024.pdf",                 61,  "ERg4NeFLkeBO"),
    ("Quy chế CTSV ĐHBK Hà Nội 2025.3.10_final.pdf",          68,  "ESCBPVQlzNFOlglsvRwxAZYBqxdZc6QR"),
    ("QĐ Ban hành hướng dân triển khai chính sachsHT cho SV khuyết tật.pdf", 68, "EQqI98V8izpJ"),
    ("Mẫu-Giấy xác nhận công nợ - Thụy.doc",                 100, "IQALcDMWqSVw"),
]


# ===========================================================================
# Local file extraction (trích xuất text từ file .doc/.docx/.pdf)
# ===========================================================================

def _extract_pdf(path: str) -> str:
    """Trích xuất text từ file PDF."""
    if _PdfReader is None:
        return ""
    try:
        reader = _PdfReader(path)
        pages = [page.extract_text() or "" for page in reader.pages]
        return clean_text("\n".join(pages))
    except Exception:
        return ""


def _extract_docx(path: str) -> str:
    """Trích xuất text từ file .docx."""
    if not _HAS_DOCX:
        return ""
    try:
        doc = _docx.Document(path)
        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    parts.append(" | ".join(row_text))
        return clean_text("\n".join(parts))
    except Exception:
        return ""


def _extract_doc_binary(path: str) -> str:
    """Trích xuất text từ file .doc cũ (OLE2) bằng COM hoặc binary fallback."""
    if _HAS_WIN32:
        try:
            _pythoncom.CoInitialize()
            word = _win32com.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            doc = word.Documents.Open(os.path.abspath(path), False, True)
            text = doc.Content.Text
            doc.Close(False)
            word.Quit()
            text = text.replace("\r", "\n")
            text = re.sub(r"\n{3,}", "\n\n", text)
            return clean_text(text)
        except Exception as e:
            logger.debug(f"COM extract failed for {path}: {e}")

    # Binary fallback
    try:
        with open(path, "rb") as f:
            content = f.read()
        pattern = b"(?:[\x00-\xff][\x00-\x05]){3,}"
        matches = re.findall(pattern, content)
        pieces: list[str] = []
        for m in matches:
            try:
                s = m.decode("utf-16-le", errors="ignore")
                s = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]+", " ", s).strip()
                if len(s) >= 3 and not all(c in " \t\r\n" for c in s):
                    pieces.append(s)
            except Exception:
                pass
        full = "\n".join(pieces)
        full = re.sub(r"[ \t]+", " ", full)
        full = re.sub(r"\n\s*\n", "\n", full)
        return clean_text(full)
    except Exception:
        return ""


def _get_file_text(fpath: str) -> str:
    """Dispatch đọc file theo extension.
    
    Nếu là file .doc, ưu tiên kiểm tra xem đã có file .docx tương ứng chưa.
    """
    ext = os.path.splitext(fpath)[1].lower()
    if ext == ".docx":
        return _extract_docx(fpath)
    if ext == ".pdf":
        return _extract_pdf(fpath)
    if ext == ".doc":
        docx_path = os.path.splitext(fpath)[0] + ".docx"
        if os.path.exists(docx_path):
            text = _extract_docx(docx_path)
            if text:
                return text
        text = _extract_docx(fpath)
        return text if text else _extract_doc_binary(fpath)
    return ""


def load_local_attachment_data(data_fetch_dir: Path) -> dict[str, dict]:
    """
    Đọc toàn bộ file từ data_fetch/ theo FILE_MAPPING.
    Trả về dict: {signature: {filename, doc_id, text}}.
    """
    attachment_data: dict[str, dict] = {}
    if not data_fetch_dir.exists():
        logger.warning(f"Thư mục data_fetch không tồn tại: {data_fetch_dir}")
        return attachment_data

    # Tự động convert .doc -> .docx tùy theo hệ điều hành (Windows / Linux)
    try:
        run_convert_script(str(data_fetch_dir))
    except Exception as e:
        logger.warning(f"Lỗi khi thực thi script convert_docs: {e}")

    found = 0
    for fname, doc_id, sig in FILE_MAPPING:
        fpath = str(data_fetch_dir / fname)
        if os.path.exists(fpath):
            text = _get_file_text(fpath)
            attachment_data[sig] = {
                "filename": fname,
                "doc_id": doc_id,
                "text": text,
            }
            found += 1
            logger.debug(f"  Đọc file cục bộ: {fname} ({'có text' if text else 'rỗng'})")
        else:
            logger.debug(f"  Không tìm thấy file: {fname}")

    logger.info(f"Đọc được {found}/{len(FILE_MAPPING)} file từ {data_fetch_dir}")
    return attachment_data


def enrich_article_with_local_files(
    article: dict[str, Any],
    attachment_data: dict[str, dict],
) -> dict[str, Any]:
    """
    Enrich bài viết với nội dung file cục bộ từ data_fetch/.
    Cập nhật:
      - link['content'], link['fetch_status'], link['local_filename']
      - article['rag_text'] với nội dung file đính kèm
    """
    if not attachment_data:
        return article

    rag_text: str = article.get("rag_text", "")

    for link in article.get("links", []):
        url = link.get("url", "")
        for sig, item in attachment_data.items():
            if sig in url:
                link["content"] = item["text"]
                link["fetch_status"] = "ok" if item["text"] else "empty"
                link["local_filename"] = item["filename"]

                if item["text"]:
                    # Rewrite anchor text trong rag_text để LLM rõ đây là biểu mẫu
                    old_link = f"[{link['anchor_text']}]({url})"
                    new_link = f"[Biểu mẫu đính kèm: {item['filename']}]({url})"
                    if old_link in rag_text:
                        rag_text = rag_text.replace(old_link, new_link)
                    else:
                        escaped_url = re.escape(url)
                        rag_text = re.sub(
                            rf'\[([^\]]+)\]\({escaped_url}\)',
                            rf'[Biểu mẫu đính kèm: {item["filename"]}]({url})',
                            rag_text,
                        )
                break  # Mỗi link chỉ match 1 file

    article["rag_text"] = clean_text(rag_text)
    return article


# ===========================================================================
# Pipeline chính
# ===========================================================================

def run_pipeline(
    *,
    use_detail_endpoint: bool = True,
    delay_seconds: float = 0.15,
    fetch_attachments: bool = False,
    enrich_local: bool = False,
    data_fetch_dir: Path = DEFAULT_DATA_FETCH_DIR,
    verbose: bool = False,
) -> dict[str, Any]:
    """
    Chạy toàn bộ pipeline: crawl API → clean → chunk → lưu PostgreSQL.

    Returns:
        summary dict với thống kê pipeline run.
    """
    started_at = datetime.now(timezone.utc).isoformat()
    logger.info("=" * 60)
    logger.info("HUST RAG Pipeline bắt đầu")
    logger.info(f"  use_detail_endpoint : {use_detail_endpoint}")
    logger.info(f"  fetch_attachments   : {fetch_attachments}")
    logger.info(f"  enrich_local        : {enrich_local}")
    logger.info(f"  delay               : {delay_seconds}s")
    logger.info("=" * 60)

    # ── Kết nối DB và khởi tạo schema ───────────────────────────────────────
    conn = None
    try:
        from db.connection import get_connection
        conn = get_connection()
        logger.info("Kết nối PostgreSQL thành công.")
        init_schema(conn)

        run_id = create_pipeline_run(
            conn,
            fetch_attachments=fetch_attachments,
            use_detail_endpoint=use_detail_endpoint,
            enrich_local=enrich_local,
        )
        logger.info(f"Pipeline run ID: {run_id}")
    except Exception as e:
        if conn:
            conn.close()
        raise

    try:
        # ── Load file cục bộ từ data_fetch/ ─────────────────────────────────
        attachment_data: dict[str, dict] = {}
        if enrich_local:
            logger.info(f"Đọc file cục bộ từ: {data_fetch_dir.resolve()}")
            attachment_data = load_local_attachment_data(data_fetch_dir)

        # ── Gọi API lấy danh sách bài viết ──────────────────────────────────
        logger.info("Đang gọi API GetWebTitleLst...")
        listing = post_json("HWAdmin/GetWebTitleLst", {})
        if listing.get("RespCode") != 0:
            raise RuntimeError(
                f"API trả về lỗi RespCode={listing.get('RespCode')}: "
                f"{listing.get('RespText')}"
            )

        raw_items: list[dict] = listing.get("WebTitleLst") or []
        handbook_items = [
            item for item in raw_items
            if item.get("TypeDoc") in TYPE_DOC_TO_CATEGORY
            and item.get("Status") == 1
        ]
        handbook_items.sort(
            key=lambda x: str(x.get("TimeCreate") or ""), reverse=True
        )

        total_api = len(raw_items)
        total_handbook = len(handbook_items)
        logger.info(f"API trả về {total_api} items, lọc được {total_handbook} bài Sổ tay SV.")

        # ── Xử lý từng bài viết ─────────────────────────────────────────────
        total_articles = 0
        total_chunks = 0
        attachment_stats: dict[str, int] = {}


        for idx, raw in enumerate(handbook_items, start=1):
            doc_id = raw.get("DocumentID", "?")

            # 1. Gọi API chi tiết (nếu bật)
            detail_item = None
            if use_detail_endpoint:
                try:
                    detail_resp = post_json(
                        "HWAdmin/GetWebTitleInfo", {"DocumentID": doc_id}
                    )
                    if detail_resp.get("RespCode") == 0:
                        detail_item = detail_resp.get("WebTitleInfo")
                    else:
                        logger.warning(
                            f"[{idx}/{total_handbook}] GetWebTitleInfo lỗi "
                            f"RespCode={detail_resp.get('RespCode')} cho DocumentID={doc_id}"
                        )
                except Exception as e:
                    logger.warning(f"[{idx}/{total_handbook}] Lỗi gọi detail API cho {doc_id}: {e}")

                if delay_seconds > 0 and idx < total_handbook:
                    time.sleep(delay_seconds)

            # 2. Normalize (làm sạch HTML → text + links)
            try:
                article = normalize_article(
                    raw,
                    detail_item,
                    fetch_attachments=fetch_attachments,
                    verbose=verbose,
                )
            except Exception as e:
                logger.error(f"[{idx}/{total_handbook}] Lỗi normalize bài {doc_id}: {e}")
                continue

            # Cập nhật attachment stats
            for link in article.get("links", []):
                s = link.get("fetch_status", "skipped")
                attachment_stats[s] = attachment_stats.get(s, 0) + 1

            # 3. Enrich với file cục bộ từ data_fetch/
            if enrich_local and attachment_data:
                article = enrich_article_with_local_files(article, attachment_data)

            # 3.5. Build attachments list
            attachments = []
            for i, link in enumerate(article.get("links", [])):
                if link.get("fetch_status") == "ok" and link.get("content"):
                    url = link.get("url", "")
                    content = link.get("content", "")
                    att_id = hashlib.sha256(url.encode("utf-8")).hexdigest()[:16]
                    content_hash = hashlib.sha256(content.encode("utf-8")).hexdigest()
                    attachments.append({
                        "attachment_id": att_id,
                        "name": link.get("local_filename") or link.get("anchor_text") or "Tài liệu đính kèm",
                        "url": url,
                        "local_filename": link.get("local_filename", ""),
                        "link_index": i,
                        "fetch_status": "ok",
                        "content": content,
                        "content_hash": content_hash,
                    })
            article["attachments"] = attachments

            title = article.get("title", doc_id)

            # 4. Build RAG document
            rag_doc = to_rag_document(article)

            # 5. Chunk
            chunks = build_chunks(rag_doc)

            # 6. Lưu vào PostgreSQL
            try:
                upsert_article(conn, article)
                
                # Xóa các chunk cũ của bài viết này (nếu có từ lần chạy trước mà không còn trong bộ chunks mới)
                parent_id = f"hust_sotay_{doc_id}"
                if chunks:
                    parent_id = chunks[0].get("parent_id", parent_id)
                    valid_ids = [c["id"] for c in chunks]
                    with conn.cursor() as cur:
                        cur.execute(
                            "DELETE FROM rag_chunks WHERE parent_id = %s AND id != ALL(%s)",
                            (parent_id, valid_ids)
                        )
                else:
                    with conn.cursor() as cur:
                        cur.execute("DELETE FROM rag_chunks WHERE parent_id = %s", (parent_id,))

                upsert_chunks(conn, chunks)
                conn.commit()
            except Exception as e:
                conn.rollback()
                logger.error(f"[{idx}/{total_handbook}] Lỗi lưu DB bài {doc_id}: {e}")
                continue

            total_articles += 1
            total_chunks += len(chunks)

            if verbose or idx % 10 == 0 or idx == total_handbook:
                logger.info(
                    f"[{idx:3d}/{total_handbook}] ✓ {title[:50]!r} "
                    f"→ {len(chunks)} chunks"
                )

        # ── Hoàn thành pipeline run ──────────────────────────────────────────
        finish_pipeline_run(
            conn,
            run_id,
            total_articles=total_articles,
            total_chunks=total_chunks,
            total_api_items=total_api,
            attachment_stats=attachment_stats,
            status="success",
        )

    except Exception as exc:
        # Ghi log lỗi vào DB nếu có thể
        try:
            if conn and run_id:
                finish_pipeline_run(
                    conn,
                    run_id,
                    total_articles=total_articles if "total_articles" in dir() else 0,
                    total_chunks=total_chunks if "total_chunks" in dir() else 0,
                    total_api_items=total_api if "total_api" in dir() else 0,
                    attachment_stats=attachment_stats if "attachment_stats" in dir() else {},
                    status="error",
                    error_message=str(exc),
                )
        except Exception:
            pass
        raise
    finally:
        if conn:
            conn.close()

    summary = {
        "run_id": run_id,
        "started_at": started_at,
        "finished_at": datetime.now(timezone.utc).isoformat(),
        "total_api_items": total_api,
        "total_articles": total_articles,
        "total_chunks": total_chunks,
        "fetch_attachments": fetch_attachments,
        "enrich_local": enrich_local,
        "attachment_stats": attachment_stats,
    }

    # ── In tóm tắt ───────────────────────────────────────────────────────────
    logger.info("")
    logger.info("=" * 60)
    logger.info("✅ Pipeline hoàn tất!")
    logger.info(f"   Pipeline Run ID  : {run_id}")
    logger.info(f"   Tổng items API   : {total_api}")
    logger.info(f"   Bài viết đã lưu  : {total_articles}")
    logger.info(f"   Chunks đã lưu    : {total_chunks}")
    if attachment_stats:
        logger.info(f"   Attachment stats : {attachment_stats}")
    logger.info("=" * 60)

    return summary


# ===========================================================================
# CLI
# ===========================================================================

def parse_args(argv: list[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="HUST RAG Pipeline: Crawl → Clean → Chunk → PostgreSQL",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Ví dụ:
  python pipeline.py
      Crawl đầy đủ (có gọi GetWebTitleInfo cho từng bài)

  python pipeline.py --no-detail
      Chỉ dùng dữ liệu từ GetWebTitleLst (nhanh hơn, ít chi tiết hơn)

  python pipeline.py --fetch-attachments
      Crawl + tải Google Doc / Google Drive / PDF đính kèm

  python pipeline.py --enrich-local
      Enrich với file .doc/.docx/.pdf trong thư mục data_fetch/

  python pipeline.py --enrich-local --data-fetch-dir ./data_fetch

  python pipeline.py --init-schema-only
      Chỉ khởi tạo schema PostgreSQL, không crawl

Cấu hình PostgreSQL (file .env hoặc biến môi trường):
  POSTGRES_HOST, POSTGRES_PORT, POSTGRES_DB, POSTGRES_USER, POSTGRES_PASSWORD
""",
    )
    parser.add_argument(
        "--no-detail",
        action="store_true",
        help="Bỏ qua GetWebTitleInfo (nhanh hơn nhưng ít chi tiết).",
    )
    parser.add_argument(
        "--delay",
        type=float,
        default=0.15,
        help="Delay giữa các lần gọi detail API (giây, mặc định: 0.15).",
    )
    parser.add_argument(
        "--fetch-attachments",
        action="store_true",
        help="Tải và trích xuất text từ Google Doc/Drive/PDF đính kèm.",
    )
    parser.add_argument(
        "--enrich-local",
        action="store_true",
        help="Enrich bài viết với file cục bộ từ thư mục data_fetch/.",
    )
    parser.add_argument(
        "--data-fetch-dir",
        type=Path,
        default=DEFAULT_DATA_FETCH_DIR,
        help=f"Thư mục chứa file cục bộ khi dùng --enrich-local (mặc định: {DEFAULT_DATA_FETCH_DIR}).",
    )
    parser.add_argument(
        "--init-schema-only",
        action="store_true",
        help="Chỉ khởi tạo schema PostgreSQL, không crawl dữ liệu.",
    )
    parser.add_argument(
        "--verbose",
        action="store_true",
        help="In chi tiết tiến độ từng bài và từng link.",
    )
    return parser.parse_args(argv)


def main(argv: list[str]) -> int:
    args = parse_args(argv)

    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)

    # Chế độ chỉ khởi tạo schema
    if args.init_schema_only:
        logger.info("Chế độ --init-schema-only: Khởi tạo schema PostgreSQL...")
        try:
            from db.connection import get_managed_connection, init_schema
            with get_managed_connection() as conn:
                init_schema(conn)
            logger.info("✅ Schema khởi tạo thành công.")
            return 0
        except Exception as e:
            logger.error(f"❌ Lỗi: {e}")
            return 1

    try:
        summary = run_pipeline(
            use_detail_endpoint=not args.no_detail,
            delay_seconds=max(args.delay, 0.0),
            fetch_attachments=args.fetch_attachments,
            enrich_local=args.enrich_local,
            data_fetch_dir=args.data_fetch_dir,
            verbose=args.verbose,
        )
        print(json.dumps(summary, ensure_ascii=False, indent=2))
        return 0
    except ConnectionError as e:
        logger.error(f"❌ Lỗi kết nối database: {e}")
        logger.error("   Hãy kiểm tra file .env và đảm bảo PostgreSQL đang chạy.")
        return 1
    except Exception as e:
        logger.error(f"❌ Pipeline thất bại: {e}")
        if args.verbose:
            import traceback
            traceback.print_exc()
        return 1


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
