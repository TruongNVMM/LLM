import os
import json
import re
import sys
import io
import unicodedata
from pathlib import Path
from pypdf import PdfReader
import docx

try:
    import win32com.client
    import pythoncom
    HAS_WIN32 = True
except ImportError:
    HAS_WIN32 = False

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

def clean_emojis_and_symbols(value: str) -> str:
    if not value:
        return ""
    cleaned = []
    for ch in value:
        cat = unicodedata.category(ch)
        # Strip control characters (Cc) except newline, carriage return, and tab
        if cat == "Cc" and ch not in ("\n", "\r", "\t"):
            continue
        if cat == "So" or ch in ("□", "☐", "■", "▪", "▫", "♦", "●", "○", "★", "☆", "▶", "►", "◄", "▼", "▲"):
            continue
        cleaned.append(ch)
    res = "".join(cleaned)
    res = re.sub(r' +', ' ', res)
    return res.strip()

def extract_pdf(path):
    try:
        reader = PdfReader(path)
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages).strip()
    except Exception:
        return ""

def extract_docx(path):
    try:
        doc = docx.Document(path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
        return "\n".join(full_text).strip()
    except Exception:
        return ""

def extract_doc_binary(path):
    if HAS_WIN32:
        try:
            pythoncom.CoInitialize()
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            doc = word.Documents.Open(os.path.abspath(path), False, True)
            text = doc.Content.Text
            doc.Close(False)
            word.Quit()
            
            # Clean up word COM special chars
            text = text.replace('\r', '\n')
            text = re.sub(r'\n{3,}', '\n\n', text)
            return text.strip()
        except Exception as e:
            print(f"COM extract failed for {path}: {e}")
            pass

    # Fallback to binary parsing
    try:
        with open(path, 'rb') as f:
            content = f.read()
        pattern = b'(?:[\x00-\xff][\x00-\x05]){3,}'
        matches = re.findall(pattern, content)
        text_pieces = []
        for m in matches:
            try:
                s = m.decode('utf-16-le', errors='ignore')
                s = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]+', ' ', s).strip()
                if len(s) >= 3 and not all(c in ' \t\r\n' for c in s):
                    text_pieces.append(s)
            except Exception:
                pass
        full = "\n".join(text_pieces)
        full = re.sub(r'[ \t]+', ' ', full)
        full = re.sub(r'\n\s*\n', '\n', full)
        return full.strip()
    except Exception:
        return ""

def get_file_text(fpath):
    ext = os.path.splitext(fpath)[1].lower()
    if ext == '.docx':
        return extract_docx(fpath)
    elif ext == '.pdf':
        return extract_pdf(fpath)
    elif ext == '.doc':
        res = extract_docx(fpath)
        if res: return res
        return extract_doc_binary(fpath)
    return ""

file_mapping = [
    ("16. Đơn công nhận HP tương đương.doc", 69, "EUZMK0oVAZpB"),
    ("20. Miễn ngoại ngữ.docx", 69, "EVM4qdnP_45J"),
    ("HƯỚNG DẪN VỀ VIỆC YÊU CẦU NHẬP ĐIỂM ĐÚNG THỜI HẠN.docx", 69, "EUDPbbk1m01J"),
    ("08.2 Đơn thắc mắc ĐATN.doc", 69, "EdDCJnCWLlRH"),
    ("22. Mẫu đơn rút học phần.doc", 69, "Eap5qh9aMf5E"),
    ("03.4 Đơn xin thôi học_ĐT.QT06.BM.01_Rev 16.4.2023.doc", 69, "EdHS7TFGfW1G"),
    ("04.3 Đơn xin nghỉ dài hạn_06.02 ĐT.QT06.BM.02.doc", 69, "EVDUxHPtC3NP"),
    ("17.1 Đơn xin chuyển trường ĐH-Rev 01_Updated 13.4.2023.doc", 69, "ETUYBfTNy3pA"),
    ("13. Mẫu đơn đăng ký sang hệ VHVL.docx", 69, "EUyvO31qrjRD"),
    ("19. Trích sao bảng điểm.docx", 69, "EeCW4MRDLJZN"),
    ("Hướng dẫn hoãn thi-thi bù.docx", 69, "EVmtZeRjgWFH"),
    ("15.1 Đơn hoãn thi_07.01 DT.QT.07.BM.01_v2.doc", 69, "ESj7AYgEYQZM"),
    ("Đơn xin chuyển ngành học KSCS.doc", 69, "ERR1zxfXQIFO"),
    ("20230710 1. QĐ Học bổng KKHT 2023.pdf", 61, "ESCBPVQlzNFOlglsvRwxAZYBT5EgEPuy"),
    ("Quy định Học bổng Trần Đại Nghĩa 2025.docx", 61, "ERczoBK8xr9P"),
    ("Quy định xét cấp HB tài trợ 2024 LasVer.pdf", 61, "EbEKhfFyCe9C"),
    ("QĐ HB gắn kết quê hương 02.4.2024.pdf", 61, "ERg4NeFLkeBO"),
    ("Quy chế CTSV ĐHBK Hà Nội 2025.3.10_final.pdf", 68, "ESCBPVQlzNFOlglsvRwxAZYBqxdZc6QR"),
    ("QĐ Ban hành hướng dân triển khai chính sachsHT cho SV khuyết tật.pdf", 68, "EQqI98V8izpJ"),
    ("Mẫu-Giấy xác nhận công nợ - Thụy.doc", 100, "IQALcDMWqSVw"),
]

# Extract content for all mapped files
attachment_data = {}
for fname, doc_id, sig in file_mapping:
    fpath = os.path.join("data_fetch", fname)
    if os.path.exists(fpath):
        text = get_file_text(fpath)
        attachment_data[sig] = {
            "filename": fname,
            "doc_id": doc_id,
            "text": text
        }

# Update sotay_articles.json
with open('data/hust_sotay/processed/sotay_articles.json', 'r', encoding='utf-8') as f:
    articles = json.load(f)

updated_links_count = 0
for article in articles:
    attached_sections = []
    rag_text = article.get('rag_text', '')
    
    for link in article.get('links', []):
        url = link.get('url', '')
        for sig, item in attachment_data.items():
            if sig in url:
                link['content'] = item['text']
                link['fetch_status'] = 'ok' if item['text'] else 'empty'
                link['local_filename'] = item['filename']
                updated_links_count += 1
                
                if item['text']:
                    # Build attachment section
                    attached_sections.append(f"[Nội dung tài liệu đính kèm: {item['filename']}]\n{item['text']}\n[/Nội dung]")
                    
                    # Rewrite the anchor text in the RAG text so the LLM explicitly knows it's a form!
                    # E.g. [Tại đây](url) -> [Biểu mẫu: Mẫu đơn xin thôi học.doc](url)
                    old_link_str = f"[{link['anchor_text']}]({url})"
                    new_link_str = f"[Biểu mẫu đính kèm: {item['filename']}]({url})"
                    
                    if old_link_str in rag_text:
                        rag_text = rag_text.replace(old_link_str, new_link_str)
                    else:
                        # Sometimes markdownify escapes brackets or changes spaces
                        # Let's do a fallback regex replace for just the URL
                        # E.g. [Tại đây](url) or [TẠI ĐÂY](url)
                        # We use re.sub to replace any [text](url) with our explicit form name
                        escaped_url = re.escape(url)
                        rag_text = re.sub(rf'\[([^\]]+)\]\({escaped_url}\)', rf'[Biểu mẫu đính kèm: {item['filename']}]({url})', rag_text)

    # Update rag_text with newly fetched local attachments and link renames
    if attached_sections:
        extra_text = "\n\n" + "\n\n".join(attached_sections)
        if extra_text not in rag_text:
            rag_text += extra_text
            
    article['rag_text'] = clean_emojis_and_symbols(rag_text)

# Write back updated sotay_articles.json
with open('data/hust_sotay/processed/sotay_articles.json', 'w', encoding='utf-8') as f:
    json.dump(articles, f, ensure_ascii=False, indent=2)

# Write back updated sotay_articles.jsonl
with open('data/hust_sotay/processed/sotay_articles.jsonl', 'w', encoding='utf-8') as f:
    for art in articles:
        f.write(json.dumps(art, ensure_ascii=False) + '\n')

# Write back updated rag_documents.jsonl
rag_docs = []
for article in articles:
    rag_docs.append({
        "id": f"hust_sotay_{article['doc_id']}",
        "doc_type": "hust_student_handbook_article",
        "title": article["title"],
        "text": article["rag_text"],
        "metadata": {
            "source": "HUST Sổ tay sinh viên",
            "source_url": article["source_url"],
            "doc_id": article["doc_id"],
            "type_doc": article["type_doc"],
            "category_code": article["category_code"],
            "category_name": article["category_name"],
            "time_create": article["time_create"],
            "emails": article.get("keywords", {}).get("emails", []),
            "tags": article.get("keywords", {}).get("tags", []),
            "links": article["links"],
        },
    })

with open('data/hust_sotay/processed/rag_documents.jsonl', 'w', encoding='utf-8') as f:
    for doc in rag_docs:
        f.write(json.dumps(doc, ensure_ascii=False) + '\n')

print(f"🎉 Đã tích hợp thành công {updated_links_count} file tài liệu từ data_fetch vào dữ liệu RAG!")
